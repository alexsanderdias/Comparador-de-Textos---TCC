from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from zipfile import BadZipFile

import docx
from fastapi import UploadFile
from PyPDF2 import PdfReader
from PyPDF2.errors import PdfReadError

from app.core.config import Settings


class DocumentReadError(ValueError):
    """Raised when an uploaded file cannot be processed."""


@dataclass
class UploadedDocument:
    filename: str
    text: str


def _file_extension(filename: str) -> str:
    return Path(filename).suffix.lower()


def _decode_plain_text(raw_bytes: bytes) -> str:
    for encoding in ("utf-8", "latin-1"):
        try:
            return raw_bytes.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise DocumentReadError("Não foi possível decodificar o arquivo de texto enviado.")


def _read_pdf(raw_bytes: bytes) -> str:
    try:
        reader = PdfReader(BytesIO(raw_bytes))
        parts = []
        for page in reader.pages:
            page_text = page.extract_text() or ""
            if page_text.strip():
                parts.append(page_text)
        return "\n".join(parts)
    except (PdfReadError, ValueError, TypeError, NotImplementedError) as exc:
        raise DocumentReadError(
            "Não foi possível ler o PDF enviado. Verifique se o arquivo não está "
            "corrompido, protegido ou sem texto selecionável."
        ) from exc


def _read_docx(raw_bytes: bytes) -> str:
    try:
        document = docx.Document(BytesIO(raw_bytes))
        return "\n".join(paragraph.text for paragraph in document.paragraphs if paragraph.text)
    except (BadZipFile, ValueError, TypeError, KeyError) as exc:
        raise DocumentReadError(
            "Não foi possível ler o DOCX enviado. Verifique se o arquivo não está "
            "corrompido ou incompleto."
        ) from exc


async def read_uploaded_document(
    upload_file: UploadFile,
    settings: Settings,
) -> UploadedDocument:
    filename = upload_file.filename or "arquivo_sem_nome"
    extension = _file_extension(filename)

    if extension not in settings.allowed_extensions:
        raise DocumentReadError(
            f"Formato não suportado: {extension or 'sem extensão'}. "
            f"Use: {', '.join(settings.allowed_extensions)}."
        )

    raw_bytes = await upload_file.read()
    if not raw_bytes:
        raise DocumentReadError(f"O arquivo '{filename}' foi enviado vazio.")

    if len(raw_bytes) > settings.max_file_size_bytes:
        raise DocumentReadError(
            f"O arquivo '{filename}' excede o limite de {settings.max_file_size_mb} MB."
        )

    if extension == ".txt":
        text = _decode_plain_text(raw_bytes)
    elif extension == ".pdf":
        text = _read_pdf(raw_bytes)
    elif extension == ".docx":
        text = _read_docx(raw_bytes)
    else:
        raise DocumentReadError("Formato de arquivo ainda não implementado.")

    cleaned_text = text.strip()
    if not cleaned_text:
        raise DocumentReadError(
            f"O arquivo '{filename}' não possui conteúdo textual legível."
        )

    return UploadedDocument(filename=filename, text=cleaned_text)
